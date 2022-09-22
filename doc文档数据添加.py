from docx import Document


# 创建Document对象
doc = Document()
# 添加标题
doc.add_heading('春晓', level=1)
# 添加段落
doc.add_paragraph('作者：孟浩然')
doc.add_paragraph('春眠不觉晓，')

para2 = doc.add_paragraph()
para2.add_run('处处闻啼鸟；').font.italic = True

# 向段落添加分页
doc.add_page_break()

para3 = doc.add_paragraph().add_run('夜来风雨声，')
para3.font.bold = True
para4 = doc.add_paragraph().add_run('花落知多少。')
para4.font.italic = True

doc.save('file/春晓.docx')
